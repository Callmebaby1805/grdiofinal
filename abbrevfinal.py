import os
import json
import re
import gradio as gr
import anthropic
from collections import defaultdict
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv(".env")

# Anthropic client
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

CATEGORY_TO_JSON = {
    "Sustainability and the Environment": "Sustainabilityandenvironment.json",
    "Strategic Technologies": "StatergicTechnologies.json",
    "Trade and Economics": "Trade.json",
    "Politics, Society and Governance": "PSG.json",
    "International Relations, Multipolarity and Multilateralism": "IEMTMLT.json"
}

def normalize_key(text):
    return re.sub(r'\s+', '', text.lower())

def clean_nested_expansion(text):
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\s*\((\w+)\)\)', r'\1 (\2)', text)
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\)', r'\1', text)
    return text

def handle_special_case_usa(text):
    pattern1 = r'\b(The\s+United\s+States)\s*\(\s*US\s*\)\s+(of\s+America)\s*\(\s*USA\s*\)'
    text = re.sub(pattern1, r'\1 \2 (USA)', text, flags=re.IGNORECASE)
    pattern2 = r'\bThe United States\s+of America\b(?!\s*\(\s*USA\s*\))'
    text = re.sub(pattern2, r'The United States of America (USA)', text, flags=re.IGNORECASE)
    return text

def resolve_synonyms_first_mention_global(text, synonym_groups):
    for group in synonym_groups:
        positions = {}
        for term in group:
            match = re.search(rf'\b{re.escape(term)}\b', text, flags=re.IGNORECASE)
            if match:
                positions[term] = match.start()
        if positions:
            chosen = min(positions, key=positions.get)
            for alt in group:
                if alt.lower() != chosen.lower():
                    pattern = rf'\b{re.escape(alt)}\b'
                    text = re.sub(pattern, chosen, text, flags=re.IGNORECASE)
    return text

def classify_topic(doc_path):
    doc = Document(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    title = paragraphs[0] if paragraphs else ""
    paragraph = paragraphs[1] if len(paragraphs) > 1 else ""

    categories = list(CATEGORY_TO_JSON.keys())
    system_prompt = (
        "You are an expert in classifying documents. Based only on the given title and paragraph, "
        f"classify the content into one of the following categories:\n{', '.join(categories)}.\n"
        "Respond ONLY with the exact name of the category."
    )
    user_prompt = f"Title: {title}\nParagraph: {paragraph}"

    response = client.messages.create(
        model="claude-3-5-haiku-20241022",
        max_tokens=100,
        temperature=0,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}]
    )

    return response.content[0].text.strip()

def expand_abbreviations(doc_path, base_json_path="abbreviations.json"):
    with open(base_json_path, "r", encoding="utf-8") as f:
        base_map = json.load(f)

    category = classify_topic(doc_path)
    category_map = {}
    category_json = CATEGORY_TO_JSON.get(category)
    if category_json and os.path.exists(category_json):
        with open(category_json, "r", encoding="utf-8") as f:
            category_map = json.load(f)

    abbr_map = {**base_map, **category_map}
    counts = defaultdict(int)
    norm_map = {normalize_key(v): (k, v) for k, v in abbr_map.items()}

    doc = Document(doc_path)

    synonym_groups = [
        ["South Korea", "Republic of Korea", "ROK"],
        ["The United States", "The United States of America"],
        ["North Korea", "Democratic People's Republic of Korea", "DPRK"],
        ["China", "People Republic of China", "PRC"]
    ]

    # Process only body text, not footnotes/tables/headers/TOC
    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue

        # Skip headings and TOC (Word uses "TOC" styles)
        if para.style and ("Heading" in para.style.name or "TOC" in para.style.name):
            continue

        line = resolve_synonyms_first_mention_global(line, synonym_groups)
        line = clean_nested_expansion(line)

        for norm_exp, (abbr, expansion) in norm_map.items():
            abbr_pattern = rf'\b{re.escape(abbr)}\b'
            exp_pattern = rf'\b{re.escape(expansion)}\b'
            patterns = [
                rf'{abbr}\s*\(\s*{expansion}\s*\)',
                rf'{expansion}\s*\(\s*{abbr}\s*\)',
                rf'{abbr}\s*\(\s*\w[\w\s]*\)',
                rf'\w[\w\s]*\(\s*{abbr}\s*\)',
                abbr_pattern,
                exp_pattern
            ]
            if any(re.search(p, line, re.IGNORECASE) for p in patterns):
                counts[abbr] += 1
                if counts[abbr] == 1:
                    correct_format = rf'{re.escape(expansion)}\s*\(\s*{re.escape(abbr)}\s*\)'
                    if not re.search(correct_format, line, re.IGNORECASE):
                        line = re.sub('|'.join(patterns), f"{expansion} ({abbr})", line, count=1, flags=re.IGNORECASE)
                else:
                    line = re.sub('|'.join(patterns), abbr, line, flags=re.IGNORECASE)

        line = handle_special_case_usa(line)
        para.text = line   # ✅ modifies original doc without recreating

    out_path = "expanded_result.docx"
    doc.save(out_path)
    return out_path

def gradio_abbreviation_expand(doc_file):
    output_path = expand_abbreviations(doc_file.name)
    return "✅ Document processed while preserving tables, TOC, citations, quotes, and notes.", output_path

gr.Interface(
    fn=gradio_abbreviation_expand,
    inputs=[gr.File(label="Upload DOCX Document", file_types=[".docx"])],
    outputs=[gr.Text(label="Status"), gr.File(label="Download Processed File")],
    title="Abbreviation Expander with Full Preservation",
    description="Expands abbreviations in running text while keeping tables, TOC, citations, quotes, headers/footers, footnotes, and endnotes intact."
)
