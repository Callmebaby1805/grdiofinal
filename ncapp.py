# # import tempfile
# # import zipfile
# # from lxml import etree
# # import gradio as gr

# # # import your two modules
# # import nandpercal
# # import currconv


# # def process_pipeline(docx_file):
# #     """
# #     Pipeline that runs nandpercal first, then currconv.
# #     """

# #     # Step 1: Run nandpercal process_docx
# #     temp1 = nandpercal.process_docx(docx_file)

# #     # Step 2: Run currconv process_docx on nandpercal’s output
# #     with open(temp1, "rb") as f1:
# #         class TempFile:
# #             name = temp1
# #         result = currconv.process_docx(TempFile())

# #     return result


# # # ========================
# # # Gradio UI
# # # ========================
# # iface = gr.Interface(
# #     fn=process_pipeline,
# #     inputs=gr.File(file_types=[".docx"], label="Upload Word Document"),
# #     outputs=gr.File(label="Download Modified Document"),
# #     title="Numbers + Percentages + Currency Conversion to SGD",
# #     description=(
# #         "Pipeline:\n"
# #         "1. Normalizes numbers & percentages (nandpercal).\n"
# #         "2. Converts currencies into SGD (currconv).\n\n"
# #         "Preserves document formatting, tables, footnotes, and skips text inside quotes/brackets."
# #     ),
# # )

# # if __name__ == "__main__":
# #     iface.launch()
# import tempfile
# import zipfile
# from lxml import etree
# import gradio as gr

# # import your three modules
# import abbrevfinal
# import nandpercal
# import currconv


# def process_pipeline(docx_file):
#     """
#     Pipeline that runs abbrevfinal first,
#     then nandpercal,
#     then currconv.
#     """

#     # Step 1: Run abbrevfinal expand_abbreviations
#     temp1 = abbrevfinal.expand_abbreviations(docx_file.name)

#     # Step 2: Run nandpercal process_docx on abbrevfinal’s output
#     class TempFile1:
#         name = temp1
#     temp2 = nandpercal.process_docx(TempFile1())

#     # Step 3: Run currconv process_docx on nandpercal’s output
#     class TempFile2:
#         name = temp2
#     result = currconv.process_docx(TempFile2())

#     return result


# # ========================
# # Gradio UI
# # ========================
# iface = gr.Interface(
#     fn=process_pipeline,
#     inputs=gr.File(file_types=[".docx"], label="Upload Word Document"),
#     outputs=gr.File(label="Download Modified Document"),
#     title="Abbreviations + Numbers/Percentages + Currency Conversion to SGD",
#     description=(
#         "Pipeline:\n"
#         "1. Expands abbreviations (abbrevfinal).\n"
#         "2. Normalizes numbers & percentages (nandpercal).\n"
#         "3. Converts currencies into SGD (currconv).\n\n"
#         "Preserves document formatting, tables, footnotes, and skips text inside quotes/brackets."
#     ),
# )

# if __name__ == "__main__":
#     iface.launch()


import os
import gradio as gr
from docx import Document
from docx.shared import RGBColor

# import your five modules
import abbrevfinal
import nandpercal
import currconv
import footnoteexcute
import pmptexcuter   # wraps tTesting.TextCorrector


def highlight_changes(original_path, final_path):
    """Compare original vs final docx and highlight differences in yellow in the final doc."""
    orig_doc = Document(original_path)
    final_doc = Document(final_path)

    # simple paragraph-by-paragraph comparison
    for i, (p_orig, p_final) in enumerate(zip(orig_doc.paragraphs, final_doc.paragraphs)):
        if p_orig.text.strip() != p_final.text.strip():
            # highlight the whole paragraph runs in yellow
            for run in p_final.runs:
                run.font.highlight_color = 7  # 7 = yellow in MS Word

    # save as new file
    base, ext = os.path.splitext(final_path)
    highlighted_path = base + "_highlighted" + ext
    final_doc.save(highlighted_path)
    return highlighted_path


def process_pipeline(docx_file):
    """
    Pipeline:
    1. Expand abbreviations (abbrevfinal)
    2. Normalize numbers/percentages (nandpercal)
    3. Convert currencies to SGD (currconv)
    4. Correct footnotes (footnoteexcute)
    5. Run final copyediting with TextCorrector (tTesting via pmptexcuter)
    6. Highlight differences (yellow) vs original doc
    """

    original_path = docx_file.name

    # Step 1: abbrevfinal
    temp1 = abbrevfinal.expand_abbreviations(original_path)

    # Step 2: nandpercal
    class TempFile1:
        name = temp1
    temp2 = nandpercal.process_docx(TempFile1())

    # Step 3: currconv
    class TempFile2:
        name = temp2
    temp3 = currconv.process_docx(TempFile2())

    # Step 4: footnoteexcute
    class TempFile3:
        name = temp3
    temp4 = footnoteexcute.process_docx(TempFile3())

    if not temp4 or not str(temp4).endswith(".docx") or not os.path.exists(temp4):
        raise RuntimeError(f"Footnote step failed: {temp4}")

    # Step 5: tTesting (via pmptexcuter)
    class TempFile4:
        name = temp4
    final_out, logs = pmptexcuter.process_docx(TempFile4())

    # Step 6: Highlight changes compared to original
    highlighted_out = highlight_changes(original_path, final_out)

    return highlighted_out


# ========================
# Gradio UI
# ========================
iface = gr.Interface(
    fn=process_pipeline,
    inputs=gr.File(file_types=[".docx"], label="Upload Word Document"),
    outputs=gr.File(label="Download Final Document with Changes Highlighted"),
    title="Full Pipeline: Abbreviations + Numbers + Currency + Footnotes + Copyediting + Highlight Changes",
    description=(
        "Pipeline:\n"
        "1. Expands abbreviations (abbrevfinal).\n"
        "2. Normalizes numbers & percentages (nandpercal).\n"
        "3. Converts currencies into SGD (currconv).\n"
        "4. Fixes footnotes (footnoteexcute).\n"
        "5. Runs TextCorrector from tTesting.py (via pmptexcuter).\n"
        "6. Compares original vs final document and highlights all changed text in yellow.\n\n"
        "Preserves document formatting, tables, TOC, citations, and notes."
    ),
)

if __name__ == "__main__":
    iface.launch(share=True)
